"""
Document processing and splitting logic for Manusplit.
Handles the core functionality of splitting documents by word count.
"""
import os
import logging
from pathlib import Path
from docx import Document
import utils
from typing import List, Tuple, Dict, Generator, Union


class DocumentSplitter:
    """Handles splitting documents into smaller parts based on word count."""
    
    def __init__(self, settings):
        """
        Initialize the document splitter.
        
        Args:
            settings: Settings object with configuration
        """
        self.settings = settings
        self.logger = logging.getLogger(__name__)
    
    def process_file(self, file_path: str, callback=None) -> Dict:
        """
        Process a single file, splitting it if needed.
        
        Args:
            file_path (str): Path to file to process
            callback (callable): Optional callback function for progress updates
                                Function signature: callback(status, progress, message)
        
        Returns:
            dict: Processing results with stats
        """
        result = {
            "success": False,
            "message": "",
            "original_path": file_path,
            "total_words": 0,
            "parts_created": 0,
            "output_files": []
        }
        
        # Check file access
        can_access, error_msg = utils.check_file_access(file_path)
        if not can_access:
            result["message"] = f"Error accessing file: {error_msg}"
            self.logger.error(f"Cannot access {file_path}: {error_msg}")
            return result
            
        # Determine file type and use appropriate handler
        file_path = str(file_path)  # Ensure string
        _, ext = os.path.splitext(file_path.lower())
        
        try:
            if ext == ".docx":
                return self._process_docx(file_path, callback)
            elif ext == ".txt":
                return self._process_txt(file_path, callback)
            else:
                result["message"] = f"Unsupported file type: {ext}"
                self.logger.warning(f"Unsupported file type: {ext}")
                return result
        except Exception as e:
            self.logger.exception(f"Error processing {file_path}: {str(e)}")
            result["message"] = f"Error: {str(e)}"
            return result
    
    def _process_docx(self, file_path: str, callback=None) -> Dict:
        """
        Process a DOCX file.
        
        Args:
            file_path (str): Path to DOCX file
            callback (callable): Progress callback
            
        Returns:
            dict: Processing results
        """
        result = {
            "success": False,
            "message": "",
            "original_path": file_path,
            "total_words": 0,
            "parts_created": 0,
            "output_files": []
        }
        
        try:
            # Load document
            if callback:
                callback("loading", 0, f"Loading {os.path.basename(file_path)}...")
                
            doc = Document(file_path)
            
            # Count total words
            paragraphs = list(self._get_docx_paragraphs(doc))
            total_words = sum(utils.count_words(p.text) for p in paragraphs)
            result["total_words"] = total_words
            
            # Check if splitting is needed
            max_words = self.settings.get("max_words")
            skip_under_limit = self.settings.get("skip_under_limit")
            
            if total_words <= max_words and skip_under_limit:
                msg = f"Skipped: {os.path.basename(file_path)} (only {utils.format_word_count(total_words)} words)"
                result["message"] = msg
                result["success"] = True
                if callback:
                    callback("skipped", 100, msg)
                return result
            
            # Prepare for splitting
            output_folder = self.settings.get("output_folder")
            preserve_formatting = self.settings.get("preserve_formatting")
            
            # Start splitting
            part_num = 1
            current_doc = Document()
            current_words = 0
            progress_count = 0
            
            for i, para in enumerate(paragraphs):
                # Calculate progress
                progress = min(100, int((i / len(paragraphs)) * 100))
                if callback and progress > progress_count:
                    progress_count = progress
                    callback("processing", progress, f"Processing paragraph {i+1} of {len(paragraphs)}...")
                
                # Get word count for this paragraph
                para_words = utils.count_words(para.text)
                
                # Check if adding this paragraph would exceed the limit
                if current_words + para_words > max_words and current_words > 0:
                    # Save current document
                    output_path = utils.get_output_filename(file_path, part_num, output_folder)
                    if callback:
                        callback("saving", progress, f"Saving part {part_num}...")
                    
                    current_doc.save(output_path)
                    result["output_files"].append(str(output_path))
                    
                    # Start a new document
                    part_num += 1
                    current_doc = Document()
                    current_words = 0
                
                # Add paragraph to current document
                if preserve_formatting:
                    # Add with as much formatting as possible
                    p = current_doc.add_paragraph()
                    # Try to preserve basic formatting (bold, italic, etc.)
                    for run in para.runs:
                        r = p.add_run(run.text)
                        r.bold = run.bold
                        r.italic = run.italic
                        r.underline = run.underline
                else:
                    # Add as plain text
                    current_doc.add_paragraph(para.text)
                
                current_words += para_words
            
            # Save the last part
            if current_words > 0:
                output_path = utils.get_output_filename(file_path, part_num, output_folder)
                if callback:
                    callback("saving", 100, f"Saving part {part_num}...")
                
                current_doc.save(output_path)
                result["output_files"].append(str(output_path))
            
            # Update result
            result["parts_created"] = part_num
            result["success"] = True
            result["message"] = f"Split into {part_num} parts ({utils.format_word_count(total_words)} words)"
            
            if callback:
                callback("complete", 100, result["message"])
                
            return result
                
        except Exception as e:
            self.logger.exception(f"Error processing DOCX {file_path}: {str(e)}")
            result["message"] = f"Error: {str(e)}"
            if callback:
                callback("error", 100, f"Error: {str(e)}")
            return result
    
    def _process_txt(self, file_path: str, callback=None) -> Dict:
        """
        Process a TXT file.
        
        Args:
            file_path (str): Path to TXT file
            callback (callable): Progress callback
            
        Returns:
            dict: Processing results
        """
        result = {
            "success": False,
            "message": "",
            "original_path": file_path,
            "total_words": 0,
            "parts_created": 0,
            "output_files": []
        }
        
        try:
            # Load document
            if callback:
                callback("loading", 0, f"Loading {os.path.basename(file_path)}...")
                
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
                
            # Split into paragraphs (by double newline or single newline)
            paragraphs = content.split('\n\n')
            if len(paragraphs) == 1:  # If no double newlines, try single newlines
                paragraphs = content.split('\n')
            
            # Remove empty paragraphs
            paragraphs = [p.strip() for p in paragraphs if p.strip()]
            
            # Count total words
            total_words = sum(utils.count_words(p) for p in paragraphs)
            result["total_words"] = total_words
            
            # Check if splitting is needed
            max_words = self.settings.get("max_words")
            skip_under_limit = self.settings.get("skip_under_limit")
            
            if total_words <= max_words and skip_under_limit:
                msg = f"Skipped: {os.path.basename(file_path)} (only {utils.format_word_count(total_words)} words)"
                result["message"] = msg
                result["success"] = True
                if callback:
                    callback("skipped", 100, msg)
                return result
            
            # Prepare for splitting
            output_folder = self.settings.get("output_folder")
            
            # Start splitting
            part_num = 1
            current_content = []
            current_words = 0
            progress_count = 0
            
            for i, para in enumerate(paragraphs):
                # Calculate progress
                progress = min(100, int((i / len(paragraphs)) * 100))
                if callback and progress > progress_count:
                    progress_count = progress
                    callback("processing", progress, f"Processing paragraph {i+1} of {len(paragraphs)}...")
                
                # Get word count for this paragraph
                para_words = utils.count_words(para)
                
                # Check if adding this paragraph would exceed the limit
                if current_words + para_words > max_words and current_words > 0:
                    # Save current content
                    output_path = utils.get_output_filename(file_path, part_num, output_folder)
                    if callback:
                        callback("saving", progress, f"Saving part {part_num}...")
                    
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write('\n\n'.join(current_content))
                    
                    result["output_files"].append(str(output_path))
                    
                    # Start a new document
                    part_num += 1
                    current_content = []
                    current_words = 0
                
                # Add paragraph to current content
                current_content.append(para)
                current_words += para_words
            
            # Save the last part
            if current_content:
                output_path = utils.get_output_filename(file_path, part_num, output_folder)
                if callback:
                    callback("saving", 100, f"Saving part {part_num}...")
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write('\n\n'.join(current_content))
                
                result["output_files"].append(str(output_path))
            
            # Update result
            result["parts_created"] = part_num
            result["success"] = True
            result["message"] = f"Split into {part_num} parts ({utils.format_word_count(total_words)} words)"
            
            if callback:
                callback("complete", 100, result["message"])
                
            return result
                
        except Exception as e:
            self.logger.exception(f"Error processing TXT {file_path}: {str(e)}")
            result["message"] = f"Error: {str(e)}"
            if callback:
                callback("error", 100, f"Error: {str(e)}")
            return result
    
    def _get_docx_paragraphs(self, doc) -> Generator:
        """
        Extract all paragraphs from a DOCX document, including tables.
        
        Args:
            doc: Document object
            
        Yields:
            Paragraph objects
        """
        # Process regular paragraphs
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                yield para
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():  # Skip empty cells
                            yield para
